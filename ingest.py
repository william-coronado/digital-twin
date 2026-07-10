import os
import json
import time
import threading
from typing import List, Dict, Any
import pypdf
import docx
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from config_loader import load_config

config = load_config()
SOURCE_DIR = config["source_documents_dir"]
VECTOR_DB_DIR = config["vector_db_dir"]
MANIFEST_PATH = os.path.join(VECTOR_DB_DIR, "metadata_manifest.json")
COLLECTION_NAME = "digital_twin_collection"

# Ensure directories exist
os.makedirs(SOURCE_DIR, exist_ok=True)
os.makedirs(VECTOR_DB_DIR, exist_ok=True)


def get_embedding_model() -> OllamaEmbeddings:
    """Returns local embedding model wrapper."""
    config = load_config()
    return OllamaEmbeddings(model=config["embedding_model"])


_db_instance = None
_db_lock = threading.RLock()  # guards instantiation and all read/write calls


def get_db() -> Chroma:
    """Single shared Chroma client for this process. Sharing one client
    (rather than one per ingest call and a separate one in QueryEngine)
    avoids two client objects concurrently hitting the same on-disk
    SQLite files from different threads (the folder watcher vs. queries)."""
    global _db_instance
    with _db_lock:
        if _db_instance is None:
            _db_instance = Chroma(
                collection_name=COLLECTION_NAME,
                embedding_function=get_embedding_model(),
                persist_directory=VECTOR_DB_DIR,
            )
        return _db_instance


def parse_file(file_path: str, rel_path: str) -> List[Document]:
    """Parses a local file (TXT, MD, PDF, DOCX, CSV, Excel, HTML) and returns a list of LangChain Documents."""
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()
    documents = []

    try:
        if ext in (".txt", ".md"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            if text.strip():
                documents.append(Document(page_content=text, metadata={"source": rel_path}))

        elif ext == ".pdf":
            reader = pypdf.PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    documents.append(
                        Document(
                            page_content=text,
                            metadata={"source": rel_path, "page": i + 1}
                        )
                    )

        elif ext == ".docx":
            doc = docx.Document(file_path)
            # Combine paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Combine tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        tables_text.append(" | ".join(cells))
            
            content = "\n".join(paragraphs)
            if tables_text:
                content += "\n\n=== Tables ===\n" + "\n".join(tables_text)
                
            if content.strip():
                documents.append(Document(page_content=content, metadata={"source": rel_path}))

        elif ext == ".csv":
            import csv
            rows_text = []
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.DictReader(f)
                for i, row in enumerate(reader):
                    row_parts = [f"{k}: {v}" for k, v in row.items() if v]
                    if row_parts:
                        rows_text.append(f"Row {i+1}: " + ", ".join(row_parts))
            content = "\n".join(rows_text)
            if content.strip():
                documents.append(Document(page_content=content, metadata={"source": rel_path}))

        elif ext == ".xlsx":
            import pandas as pd
            excel_data = pd.read_excel(file_path, sheet_name=None)
            sheets_text = []
            for sheet_name, df in excel_data.items():
                df = df.fillna("")
                rows_text = []
                for i, row in df.iterrows():
                    row_parts = [f"{col}: {val}" for col, val in row.items() if str(val).strip()]
                    if row_parts:
                        rows_text.append(f"Row {i+1}: " + ", ".join(row_parts))
                if rows_text:
                    sheets_text.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows_text))
            content = "\n\n".join(sheets_text)
            if content.strip():
                documents.append(Document(page_content=content, metadata={"source": rel_path}))

        elif ext == ".html":
            from bs4 import BeautifulSoup
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            # Remove scripts and styles
            for elem in soup(["script", "style"]):
                elem.decompose()
            text = soup.get_text(separator="\n")
            lines = (line.strip() for line in text.splitlines())
            chunks_text = (phrase.strip() for line in lines for phrase in line.split("  "))
            clean_text = "\n".join(chunk for chunk in chunks_text if chunk)
            if clean_text.strip():
                documents.append(Document(page_content=clean_text, metadata={"source": rel_path}))

    except Exception as e:
        print(f"Error parsing file {file_path}: {e}")

    return documents


def load_manifest() -> Dict[str, Any]:
    """Loads indexing manifest metadata."""
    if os.path.exists(MANIFEST_PATH):
        try:
            with open(MANIFEST_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"files": {}, "total_chunks": 0}


def save_manifest(manifest: Dict[str, Any]):
    """Saves indexing manifest metadata."""
    with open(MANIFEST_PATH, "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2)


def sync_vector_store(verbose: bool = True) -> Dict[str, Any]:
    """
    Scans the source directory and performs incremental updates to the Chroma DB.
    Detects deleted, new, and modified files.
    """
    manifest = load_manifest()

    # Hold the lock for the whole sync so a concurrent query (running the
    # shared Chroma client on another thread) can't interleave with an
    # in-progress write. Sync runs quickly, so this is fine for a
    # single-user local tool.
    with _db_lock:
        db = get_db()

        # 2. Get list of files in source directory recursively
        allowed_exts = {".txt", ".md", ".pdf", ".docx", ".xlsx", ".csv", ".html"}
        current_files = {}
        for root, _, files in os.walk(SOURCE_DIR):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in allowed_exts:
                    full_path = os.path.join(root, file)
                    # Compute path relative to SOURCE_DIR
                    rel_path = os.path.relpath(full_path, SOURCE_DIR)
                    try:
                        stat = os.stat(full_path)
                        current_files[rel_path] = {
                            "path": full_path,
                            "mtime": stat.st_mtime,
                            "size": stat.st_size,
                        }
                    except Exception as e:
                        if verbose:
                            print(f"Error accessing stats for {full_path}: {e}")

        # 3. Detect and handle deletions
        deleted_files = [fname for fname in manifest["files"] if fname not in current_files]
        for fname in deleted_files:
            if verbose:
                print(f"Detected deletion of: {fname}. Removing from vector database...")
            try:
                # Chroma deletes items using metadata filter
                db.delete(where={"source": fname})
            except Exception as e:
                # If the database or collection is empty, delete can fail; we catch and proceed
                if verbose:
                    print(f"Note: Could not delete {fname} from DB (might not exist yet): {e}")
            del manifest["files"][fname]

        # 4. Detect and handle additions / modifications
        config = load_config()
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config["chunk_size"],
            chunk_overlap=config["chunk_overlap"]
        )

        for fname, info in current_files.items():
            is_new = fname not in manifest["files"]
            is_modified = False
            if not is_new:
                old_info = manifest["files"][fname]
                is_modified = info["mtime"] > old_info.get("mtime", 0) or info["size"] != old_info.get("size", 0)

            if is_new or is_modified:
                if verbose:
                    action = "indexing new file" if is_new else "re-indexing modified file"
                    print(f"Detected change: {action} '{fname}'...")

                if is_modified:
                    # Remove old chunks before re-indexing
                    try:
                        db.delete(where={"source": fname})
                    except Exception as e:
                        if verbose:
                            print(f"Note: Could not delete old chunks for {fname} from DB: {e}")

                # Parse and chunk with relative path source
                raw_docs = parse_file(info["path"], fname)
                if not raw_docs:
                    manifest["files"][fname] = {
                        "mtime": info["mtime"],
                        "size": info["size"],
                        "chunks_count": 0
                    }
                    continue

                chunks = text_splitter.split_documents(raw_docs)

                if chunks:
                    db.add_documents(chunks)
                    manifest["files"][fname] = {
                        "mtime": info["mtime"],
                        "size": info["size"],
                        "chunks_count": len(chunks)
                    }
                else:
                    manifest["files"][fname] = {
                        "mtime": info["mtime"],
                        "size": info["size"],
                        "chunks_count": 0
                    }

        # Update total chunk count
        manifest["total_chunks"] = sum(f["chunks_count"] for f in manifest["files"].values())
        save_manifest(manifest)

    if verbose:
        print(f"Sync complete. Active files: {len(manifest['files'])}, Total Chunks: {manifest['total_chunks']}")
        
    return manifest


# Watchdog Handler
class SourceFolderHandler(FileSystemEventHandler):
    def __init__(self, callback):
        self.callback = callback
        self.last_triggered = 0.0

    def on_any_event(self, event):
        # Filter directories and only trigger on documents folder changes
        if event.is_directory:
            return
        
        filename = os.path.basename(event.src_path)
        ext = os.path.splitext(filename)[1].lower()
        if ext in (".txt", ".md", ".pdf", ".docx", ".xlsx", ".csv", ".html", ".tmp", ".crdownload"):
            # Avoid rapid multiple triggers by checking elapsed time (debounce)
            now = time.time()
            if now - self.last_triggered > 2.0:
                self.last_triggered = now
                # Wait briefly for file write completion
                time.sleep(1.0)
                self.callback()


def start_folder_watcher(callback_func) -> Observer:
    """Starts a background thread watching SOURCE_DIR."""
    event_handler = SourceFolderHandler(callback_func)
    observer = Observer()
    observer.schedule(event_handler, path=SOURCE_DIR, recursive=True)
    observer.start()
    return observer


if __name__ == "__main__":
    print(f"Syncing folder: '{SOURCE_DIR}' to vector store: '{VECTOR_DB_DIR}'...")
    sync_vector_store()
